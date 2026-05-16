"""
Word Document Generator - Generate .docx files from HTML templates.
Reused and optimized from BillGeneratorUnified.
"""
import io
import logging
from typing import Dict, Any, List
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class WordGenerator:
    """Generate Word documents from HTML content rendered by Jinja."""
    
    def __init__(self):
        pass
    
    def html_to_docx(self, html_content: str, doc_name: str) -> bytes:
        """
        Converts HTML table-heavy content to a Word Document.
        Focuses on preserving the table structure used in PWD bills.
        """
        doc = Document()
        
        # Set narrow margins for landscape-like feel (0.5 inch)
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Add a title if provided
        if doc_name:
            title = doc.add_heading(doc_name, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Process all tables in the HTML
        tables = soup.find_all('table')
        if not tables:
            # Fallback to paragraph extraction if no tables found
            for p in soup.find_all(['p', 'h1', 'h2', 'h3']):
                doc.add_paragraph(p.get_text(strip=True))
        else:
            for html_table in tables:
                rows = html_table.find_all('tr')
                if not rows:
                    continue
                
                # Determine max columns
                max_cols = 0
                for row in rows:
                    cols = row.find_all(['th', 'td'])
                    max_cols = max(max_cols, len(cols))
                
                if max_cols == 0:
                    continue
                    
                word_table = doc.add_table(rows=len(rows), cols=max_cols)
                word_table.style = 'Table Grid'
                
                for row_idx, html_row in enumerate(rows):
                    cells = html_row.find_all(['th', 'td'])
                    for col_idx, cell in enumerate(cells):
                        if col_idx < max_cols:
                            word_cell = word_table.rows[row_idx].cells[col_idx]
                            word_cell.text = cell.get_text(strip=True)
                            
                            # Styling based on tag and class
                            for paragraph in word_cell.paragraphs:
                                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                                if cell.name == 'th' or 'bold' in cell.get('class', []):
                                    run.font.bold = True
                                    run.font.size = Pt(9)
                                else:
                                    run.font.size = Pt(8)
                
                # Add spacing after each table
                doc.add_paragraph()
                
        # Save to memory
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_all(self, html_docs: Dict[str, str]) -> Dict[str, bytes]:
        """Convenience method to batch generate docx from multiple HTML strings."""
        results = {}
        for name, html in html_docs.items():
            try:
                results[name] = self.html_to_docx(html, name)
            except Exception as e:
                logger.error(f"Failed to generate Word doc for {name}: {e}")
        return results
